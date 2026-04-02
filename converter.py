"""
converter.py — GPU-accelerated PDF ↔ Word converter
Chạy trên NVIDIA Brev với GPU (H100/A100/RTX).

PDF   → Word : dùng marker-pdf (surya OCR + layout detection, GPU)
Word  → PDF  : dùng LibreOffice headless (chính xác nhất cho .docx)
Image → Word : dùng surya OCR trực tiếp (GPU-accelerated)
"""

import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Literal

import torch

# ── Detect GPU ────────────────────────────────────────────────────────────────

def get_device_info() -> dict:
    if torch.cuda.is_available():
        name = torch.cuda.get_device_name(0)
        mem  = torch.cuda.get_device_properties(0).total_memory // (1024 ** 3)
        return {"device": "cuda", "name": name, "vram_gb": mem}
    return {"device": "cpu", "name": "CPU", "vram_gb": 0}


# ── PDF → Word ────────────────────────────────────────────────────────────────

def pdf_to_word(pdf_path: str, output_dir: str | None = None) -> str:
    """
    Convert PDF → .docx using marker-pdf 1.x (GPU-accelerated).
    Returns path to the output .docx file.
    """
    from marker.models import create_model_dict
    from marker.converters.pdf import PdfConverter
    from marker.config.parser import ConfigParser
    from docx import Document

    pdf_path = Path(pdf_path)
    if output_dir is None:
        output_dir = pdf_path.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Load models (cached after first load)
    models = create_model_dict()

    # Config: tận dụng GPU Brev
    config = ConfigParser({
        "langs": "vi,en",
        "batch_multiplier": 4,
    })

    # Convert PDF → markdown
    converter = PdfConverter(
        config=config.generate_config_dict(),
        artifact_dict=models,
    )
    rendered = converter(str(pdf_path))
    full_text = rendered.markdown

    # Build .docx từ markdown output
    doc = Document()
    doc.add_heading(pdf_path.stem, level=0)

    for line in full_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    out_path = output_dir / (pdf_path.stem + ".docx")
    doc.save(str(out_path))
    return str(out_path)


# ── Word → PDF ────────────────────────────────────────────────────────────────

def word_to_pdf(docx_path: str, output_dir: str | None = None) -> str:
    """
    Convert .docx → PDF using LibreOffice headless.
    Returns path to the output .pdf file.
    """
    docx_path = Path(docx_path)
    if output_dir is None:
        output_dir = docx_path.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    libreoffice_bin = _find_libreoffice()
    if libreoffice_bin is None:
        raise RuntimeError(
            "LibreOffice không tìm thấy. Chạy: sudo apt-get install -y libreoffice"
        )

    result = subprocess.run(
        [
            libreoffice_bin,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(output_dir),
            str(docx_path),
        ],
        capture_output=True,
        text=True,
        timeout=120,
    )

    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice error:\n{result.stderr}")

    out_path = output_dir / (docx_path.stem + ".pdf")
    if not out_path.exists():
        raise FileNotFoundError(f"LibreOffice không tạo được file: {out_path}")

    return str(out_path)


# ── Image → Word ─────────────────────────────────────────────────────────────

def image_to_word(img_path: str, output_dir: str | None = None) -> str:
    """
    Convert image (JPG/PNG/TIFF/WebP/BMP) → .docx using surya OCR (GPU-accelerated).
    Returns path to the output .docx file.
    """
    from PIL import Image
    from surya.ocr import run_ocr
    from surya.model.detection.model import load_model as load_det_model
    from surya.model.detection.processor import load_processor as load_det_processor
    from surya.model.recognition.model import load_model as load_rec_model
    from surya.model.recognition.processor import load_processor as load_rec_processor
    from docx import Document

    img_path = Path(img_path)
    if output_dir is None:
        output_dir = img_path.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Mở ảnh
    image = Image.open(str(img_path)).convert("RGB")

    # Load surya models (GPU nếu có)
    det_processor = load_det_processor()
    det_model     = load_det_model()
    rec_model     = load_rec_model()
    rec_processor = load_rec_processor()

    # Chạy OCR — GPU-accelerated
    langs       = [["vi", "en"]]
    predictions = run_ocr(
        [image],
        langs,
        det_model,
        det_processor,
        rec_model,
        rec_processor,
    )

    # Build .docx từ kết quả OCR
    doc = Document()
    doc.add_heading(img_path.stem, level=0)

    for page in predictions:
        for line in page.text_lines:
            text = line.text.strip()
            if text:
                doc.add_paragraph(text)

    out_path = output_dir / (img_path.stem + ".docx")
    doc.save(str(out_path))
    return str(out_path)


def get_image_info(img_path: str) -> dict:
    """Trả về thông tin cơ bản của ảnh: width, height, format, mode."""
    from PIL import Image
    with Image.open(img_path) as im:
        return {
            "width":  im.width,
            "height": im.height,
            "format": im.format or Path(img_path).suffix.lstrip(".").upper(),
            "mode":   im.mode,
        }


# ── Helpers ───────────────────────────────────────────────────────────────────

def _find_libreoffice() -> str | None:
    for candidate in ["libreoffice", "soffice", "/usr/bin/libreoffice", "/usr/bin/soffice"]:
        if shutil.which(candidate):
            return candidate
    return None


def get_pdf_page_count(pdf_path: str) -> int:
    import fitz  # PyMuPDF
    with fitz.open(pdf_path) as doc:
        return doc.page_count


def get_docx_word_count(docx_path: str) -> int:
    from docx import Document
    doc = Document(docx_path)
    return sum(len(p.text.split()) for p in doc.paragraphs)
